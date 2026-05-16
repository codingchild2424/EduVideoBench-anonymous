#!/usr/bin/env python3
"""Convert HUMAN_EVAL_GUIDE.md to Word (.docx) format."""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


TEMPLATES_DIR = Path(__file__).resolve().parent.parent / "templates"
INPUT_MD = TEMPLATES_DIR / "HUMAN_EVAL_GUIDE.md"
OUTPUT_DOCX = TEMPLATES_DIR / "HUMAN_EVAL_GUIDE.docx"


def set_cell_shading(cell, color_hex: str):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(
        qn("w:shd"),
        {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): color_hex,
        },
    )
    shading.append(shading_elem)


def style_table(table):
    """Apply consistent styling to a table."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Style header row
    if table.rows:
        for cell in table.rows[0].cells:
            set_cell_shading(cell, "2F5496")
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.bold = True
                    run.font.size = Pt(9)
    # Style data rows
    for i, row in enumerate(table.rows[1:], 1):
        for cell in row.cells:
            if i % 2 == 0:
                set_cell_shading(cell, "D6E4F0")
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)


def parse_md_table(lines: list[str]) -> list[list[str]]:
    """Parse markdown table lines into rows of cells."""
    rows = []
    for line in lines:
        line = line.strip()
        if not line.startswith("|"):
            continue
        cells = [c.strip() for c in line.split("|")]
        # Remove empty first/last from leading/trailing |
        cells = cells[1:-1] if len(cells) > 2 else cells
        # Skip separator row (contains only dashes/colons)
        if all(re.match(r"^[\s:\-]+$", c) for c in cells):
            continue
        rows.append(cells)
    return rows


def add_table_from_md(doc, table_lines: list[str]):
    """Create a Word table from markdown table lines."""
    rows = parse_md_table(table_lines)
    if not rows:
        return
    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < n_cols:
                cell = table.cell(i, j)
                # Clean markdown formatting
                text = cell_text.replace("**", "").strip()
                cell.text = text
    style_table(table)
    doc.add_paragraph("")  # spacing after table


def clean_inline(text: str) -> list[tuple[str, bool, bool]]:
    """Parse inline markdown (bold, code) into segments.
    Returns list of (text, is_bold, is_code).
    """
    segments = []
    # Process bold and code markers
    pattern = r"(\*\*(.+?)\*\*|`(.+?)`)"
    last_end = 0
    for m in re.finditer(pattern, text):
        # Add text before match
        if m.start() > last_end:
            segments.append((text[last_end : m.start()], False, False))
        if m.group(2):  # bold
            segments.append((m.group(2), True, False))
        elif m.group(3):  # code
            segments.append((m.group(3), False, True))
        last_end = m.end()
    # Add remaining text
    if last_end < len(text):
        segments.append((text[last_end:], False, False))
    return segments if segments else [(text, False, False)]


def add_rich_paragraph(doc, text: str, style=None):
    """Add a paragraph with inline bold/code formatting."""
    p = doc.add_paragraph(style=style)
    for content, is_bold, is_code in clean_inline(text):
        run = p.add_run(content)
        if is_bold:
            run.bold = True
        if is_code:
            run.font.name = "Consolas"
            run.font.size = Pt(9)
    return p


def convert_md_to_docx(md_path: Path, docx_path: Path):
    """Convert markdown file to Word document."""
    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # -- Default font --
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)

    # Configure heading styles
    for level in range(1, 5):
        h_style = doc.styles[f"Heading {level}"]
        h_style.font.name = "Calibri"
        h_style.font.color.rgb = RGBColor(47, 84, 150)

    doc.styles["Heading 1"].font.size = Pt(18)
    doc.styles["Heading 2"].font.size = Pt(14)
    doc.styles["Heading 3"].font.size = Pt(12)
    doc.styles["Heading 4"].font.size = Pt(11)

    lines = md_path.read_text(encoding="utf-8").splitlines()

    i = 0
    in_code_block = False
    code_lines = []
    table_lines = []

    while i < len(lines):
        line = lines[i]

        # --- Code blocks ---
        if line.strip().startswith("```"):
            if in_code_block:
                # End code block
                in_code_block = False
                code_text = "\n".join(code_lines)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                # Light gray background via paragraph shading
                pPr = p._element.get_or_add_pPr()
                shading = pPr.makeelement(
                    qn("w:shd"),
                    {
                        qn("w:val"): "clear",
                        qn("w:color"): "auto",
                        qn("w:fill"): "F2F2F2",
                    },
                )
                pPr.append(shading)
                code_lines = []
                i += 1
                continue
            else:
                in_code_block = True
                code_lines = []
                i += 1
                continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # --- Tables ---
        if line.strip().startswith("|"):
            table_lines.append(line)
            i += 1
            continue
        elif table_lines:
            add_table_from_md(doc, table_lines)
            table_lines = []
            # Don't increment i, process current line

        # --- Headings ---
        if line.startswith("# ") and not line.startswith("##"):
            title = line[2:].strip()
            doc.add_heading(title, level=1)
            i += 1
            continue

        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue

        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=4)
            i += 1
            continue

        # --- Horizontal rules ---
        if line.strip() == "---":
            # Add a thin horizontal line
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pPr = p._element.get_or_add_pPr()
            pBdr = pPr.makeelement(qn("w:pBdr"), {})
            bottom = pBdr.makeelement(
                qn("w:bottom"),
                {
                    qn("w:val"): "single",
                    qn("w:sz"): "6",
                    qn("w:space"): "1",
                    qn("w:color"): "999999",
                },
            )
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # --- Blockquotes ---
        if line.strip().startswith("> "):
            text = line.strip()[2:]
            # Collect multi-line blockquote
            while i + 1 < len(lines) and lines[i + 1].strip().startswith("> "):
                i += 1
                text += " " + lines[i].strip()[2:]

            p = add_rich_paragraph(doc, text)
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            # Italic for blockquotes
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(100, 100, 100)
            i += 1
            continue

        # --- Bullet lists ---
        if re.match(r"^[-*]\s", line.strip()):
            text = line.strip()[2:]
            add_rich_paragraph(doc, text, style="List Bullet")
            i += 1
            continue

        # --- Numbered lists ---
        m = re.match(r"^(\d+)\.\s(.+)", line.strip())
        if m:
            text = m.group(2)
            add_rich_paragraph(doc, text, style="List Number")
            i += 1
            continue

        # --- Indented bullet (sub-list) ---
        if re.match(r"^\s+[-*]\s", line):
            text = line.strip()[2:]
            add_rich_paragraph(doc, text, style="List Bullet 2")
            i += 1
            continue

        # --- Empty lines ---
        if not line.strip():
            i += 1
            continue

        # --- Italic line (footer) ---
        if line.strip().startswith("*") and line.strip().endswith("*") and not line.strip().startswith("**"):
            text = line.strip().strip("*")
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)
            run.font.size = Pt(9)
            i += 1
            continue

        # --- Normal paragraph ---
        add_rich_paragraph(doc, line.strip())
        i += 1

    # Flush remaining table
    if table_lines:
        add_table_from_md(doc, table_lines)

    doc.save(str(docx_path))
    print(f"Created: {docx_path}")
    print(f"  Size: {docx_path.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    convert_md_to_docx(INPUT_MD, OUTPUT_DOCX)
